"""
Generates a scientific-style Word (.docx) report for the diploma thesis.

Sections:
  1. Uvod
  2. Metodika
     2.1 Konstrukcia datasetu
     2.2 Rozdelenie na trenovaciu, validacnu a testovaciu mnozinu
     2.3 Architektura modelu
     2.4 Vyber hyperparametrov
     2.5 Strucna funkcia
     2.6 Trenovaci postup
     2.7 Hodnotiace metriky
  3. Vysledky
     3.1 Hlavny model (IoU-balanced)
     3.2 Studia prispevku pridanych dat (data ablation)
     3.3 Experimenty s 2.5D vstupom
     3.4 Recall-optimized model (Focal-Tversky)
  4. Diskusia
  5. Zaver
  6. Pouzita literatura

Required:
    pip install python-docx

Usage:
    python export_report_word.py
"""

import argparse
import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_DIR = Path(__file__).parent
CLASS_NAMES = ["Drevo", "Kora", "Nezdrava_hrca", "Okolie", "Prasklina"]


# ═══════════════════════════════════════════════════════════════════════
# Formatting helpers
# ═══════════════════════════════════════════════════════════════════════
def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False, italic=False, size=11, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_formula(doc, formula_text, description=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(formula_text)
    run.italic = True
    run.font.size = Pt(12)
    if description:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(description)
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def shade_row(row, hex_color="D9D9D9"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)


def add_caption(doc, text):
    """Italic, smaller, centered caption under a table or figure."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def color_delta(cell, delta, threshold=1.0):
    """Color cell run green if delta>0, red if <0, when |delta|>=threshold."""
    if abs(delta) >= threshold:
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x70, 0x30) if delta > 0 \
                             else RGBColor(0xC0, 0x00, 0x00)


# ═══════════════════════════════════════════════════════════════════════
# Table builders
# ═══════════════════════════════════════════════════════════════════════
def add_dataset_summary_table(doc):
    """Tabulka 1 -- prehlad datasetov."""
    rows = [
        ["kmen1–kmen10",       "dub",                "10",   "704",
         "10 jedincov, 64 rezov/jedinec (kmen3 = 128)"],
        ["Dub_1–Dub_10",       "dub",                "10",   "643",
         "10 jedincov, 64–65 rezov/jedinec"],
        ["Dub_praskliny_a/b",  "dub (vyber)",         "—",    "100",
         "rucne vybrane, kazda obsahuje Prasklinu"],
        ["hrce_mixed",         "dub (mixed)",         "1",    "94",
         "zamerane na Nezdravu hrcu"],
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Dataset", "Druh dreva", "Pocet jedincov",
                           "Pocet snimok", "Poznamka"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    shade_row(table.rows[0], "BDD7EE")
    for r in rows:
        row = table.add_row().cells
        for i, v in enumerate(r):
            row[i].text = v
    total = table.add_row().cells
    total[0].text = "Spolu"
    total[0].paragraphs[0].runs[0].bold = True
    total[1].text = ""
    total[2].text = "—"
    total[3].text = "1541"
    total[4].text = ""
    shade_row(table.rows[-1], "E2EFDA")


def add_split_distribution_table(doc):
    """Tabulka per-class pixel distribucie pre TRAIN/VAL/TEST."""
    # Hodnoty zo split_stats.json (zaokruhlene)
    data = {
        "Test  (n=222)":  [23.23, 2.07, 0.049, 74.59, 0.058],
        "Train (n=1015)": [17.62, 1.86, 0.042, 80.46, 0.016],
        "Val   (n=254)":  [18.07, 1.94, 0.053, 79.93, 0.013],
    }
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Mnozina"] + CLASS_NAMES):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    shade_row(table.rows[0], "BDD7EE")
    for label, vals in data.items():
        row = table.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True
        for i, v in enumerate(vals):
            row[i+1].text = f"{v:.3f}%" if v < 1 else f"{v:.2f}%"


def add_sweep_table(doc, sweep_results=None):
    """Tabulka 5-konfiguracneho CE/Dice sweep-u."""
    rows = sweep_results or [
        ("0.25 / 0.75", "74.02", "33.4", 76),
        ("0.30 / 0.70", "75.19", "36.4", 98),
        ("0.35 / 0.65", "74.05", "32.3", 46),
        ("0.40 / 0.60", "74.39", "33.7", 96),
        ("0.45 / 0.55", "73.78", "31.2", 57),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Pomer CE / Dice", "mIoU (%)",
                           "IoU Prasklina (%)", "Best epocha"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    shade_row(table.rows[0], "BDD7EE")
    for r in rows:
        row = table.add_row().cells
        is_winner = (r[0] == "0.30 / 0.70")
        row[0].text = r[0] + (" *" if is_winner else "")
        row[1].text = str(r[1])
        row[2].text = str(r[2])
        row[3].text = str(r[3])
        if is_winner:
            for c in row:
                c.paragraphs[0].runs[0].bold = True
            shade_row(table.rows[-1], "E2EFDA")


def add_results_table(doc, metrics):
    per_class = metrics["per_class"]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Trieda", "IoU (%)", "Dice / F1 (%)",
                           "Presnost (%)", "Uplnost (%)", "Podpora (px)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    shade_row(table.rows[0], "BDD7EE")
    for name in CLASS_NAMES:
        m = per_class[name]
        row = table.add_row().cells
        row[0].text = name
        row[1].text = f"{m['iou']*100:.2f}"
        row[2].text = f"{m['dice']*100:.2f}"
        row[3].text = f"{m['precision']*100:.2f}"
        row[4].text = f"{m['recall']*100:.2f}"
        row[5].text = f"{m['support']:,}"
    sep = table.add_row().cells
    for c in sep:
        c.text = ""
    shade_row(table.rows[-1], "F2F2F2")
    agg = table.add_row().cells
    agg[0].text = "Priemer (macro)"
    agg[0].paragraphs[0].runs[0].bold = True
    agg[1].text = f"{metrics['mean_iou']*100:.2f}"
    agg[2].text = f"{metrics['mean_dice']*100:.2f}"
    agg[3].text = f"{metrics['mean_precision']*100:.2f}"
    agg[4].text = f"{metrics['mean_recall']*100:.2f}"
    agg[5].text = ""
    shade_row(table.rows[-1], "E2EFDA")
    pa = table.add_row().cells
    pa[0].text = "Pixelova presnost"
    pa[0].paragraphs[0].runs[0].bold = True
    pa[1].text = f"{metrics['pixel_accuracy']*100:.2f}"
    for i in range(2, 6):
        pa[i].text = ""
    shade_row(table.rows[-1], "E2EFDA")


def add_comparison_table(doc, main, ablation, label_main="v3 (plny dataset)",
                          label_other="Ablation"):
    metrics_to_show = [("IoU (%)", "iou"), ("Recall (%)", "recall"),
                       ("Precision (%)", "precision")]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Trieda", "Metrika", label_main, label_other, "Δ"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    shade_row(table.rows[0], "BDD7EE")
    for name in CLASS_NAMES:
        for j, (label, key) in enumerate(metrics_to_show):
            row = table.add_row().cells
            row[0].text = name if j == 0 else ""
            if j == 0:
                row[0].paragraphs[0].runs[0].bold = True
            row[1].text = label
            v_m = main["per_class"][name][key] * 100
            v_o = ablation["per_class"][name][key] * 100
            d = v_m - v_o
            row[2].text = f"{v_m:.2f}"
            row[3].text = f"{v_o:.2f}"
            row[4].text = f"{'+' if d >= 0 else ''}{d:.2f}"
            color_delta(row[4], d)
        sep = table.add_row().cells
        for c in sep:
            c.text = ""
        shade_row(table.rows[-1], "F8F8F8")
    for label, key in [("mIoU", "mean_iou"), ("Mean Recall", "mean_recall")]:
        agg = table.add_row().cells
        agg[0].text = "Priemer (macro)" if label == "mIoU" else ""
        if label == "mIoU":
            agg[0].paragraphs[0].runs[0].bold = True
        agg[1].text = label
        v_m = main[key] * 100
        v_o = ablation[key] * 100
        d = v_m - v_o
        agg[2].text = f"{v_m:.2f}"
        agg[3].text = f"{v_o:.2f}"
        agg[4].text = f"{'+' if d >= 0 else ''}{d:.2f}"
        color_delta(agg[4], d)
        shade_row(table.rows[-1], "E2EFDA")


def add_3way_comparison_table(doc, m1, m2, m3, labels):
    """Generic 3-way table. labels = (col_main_label, col_2_label, col_3_label, delta_label)"""
    metrics_to_show = [("IoU (%)", "iou"), ("Recall (%)", "recall"),
                       ("Precision (%)", "precision")]
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Trieda", "Metrika"] + list(labels)):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    shade_row(table.rows[0], "BDD7EE")
    for name in CLASS_NAMES:
        for j, (label, key) in enumerate(metrics_to_show):
            row = table.add_row().cells
            row[0].text = name if j == 0 else ""
            if j == 0:
                row[0].paragraphs[0].runs[0].bold = True
            row[1].text = label
            v1 = m1["per_class"][name][key] * 100
            v2 = m2["per_class"][name][key] * 100
            v3 = m3["per_class"][name][key] * 100
            d = v3 - v2
            row[2].text = f"{v1:.2f}"
            row[3].text = f"{v2:.2f}"
            row[4].text = f"{v3:.2f}"
            row[5].text = f"{'+' if d >= 0 else ''}{d:.2f}"
            color_delta(row[5], d)
        sep = table.add_row().cells
        for c in sep:
            c.text = ""
        shade_row(table.rows[-1], "F8F8F8")
    for label, key in [("mIoU", "mean_iou"), ("Mean Recall", "mean_recall")]:
        agg = table.add_row().cells
        agg[0].text = "Priemer (macro)" if label == "mIoU" else ""
        if label == "mIoU":
            agg[0].paragraphs[0].runs[0].bold = True
        agg[1].text = label
        v1 = m1[key] * 100
        v2 = m2[key] * 100
        v3 = m3[key] * 100
        d = v3 - v2
        agg[2].text = f"{v1:.2f}"
        agg[3].text = f"{v2:.2f}"
        agg[4].text = f"{v3:.2f}"
        agg[5].text = f"{'+' if d >= 0 else ''}{d:.2f}"
        color_delta(agg[5], d)
        shade_row(table.rows[-1], "E2EFDA")


# ═══════════════════════════════════════════════════════════════════════
# Section: 1. Introduction
# ═══════════════════════════════════════════════════════════════════════
def section_intro(doc):
    add_heading(doc, "1  Uvod", 1)
    add_paragraph(doc,
        "Tato praca sa zaobera semantickou segmentaciou priecnych rezov "
        "kmenov stromov ziskanych pomocou pocitacovej tomografie (CT). "
        "Cielom je automatizovana detekcia struktur dreva — najma vad "
        "(praskliny, nezdrave hrce), ktorych identifikacia ma priamy "
        "ekonomicky dopad v drevarskom priemysle. Z technickeho hladiska "
        "ide o ulohu pixelovej klasifikacie do piatich tried pri vyrazne "
        "nevyvazenom rozdeleni (vzacne triedy tvoria pod 0,1 % vsetkych "
        "pixelov)."
    )
    add_paragraph(doc,
        "Pre realizaciu bola zvolena architektura SegFormer-B2 [Xie 2021], "
        "ktora kombinuje hierarchicky transformer encoder s lahkym MLP "
        "dekoderom a dosahuje state-of-the-art vysledky na semantickej "
        "segmentacii pri umiernenych vypoctovych narocnostiach (~25 mil. "
        "parametrov). Implementacia je postavena na kniznici "
        "segmentation_models.pytorch s ImageNet predtreningom enkodera. "
        "Cely tréning aj inference prebiehaju v PyTorch s podporou AMP "
        "(automatic mixed precision)."
    )
    add_paragraph(doc, "Triedy modelu:", bold=True)
    class_desc = {
        "Drevo":         "zdrave drevo kmena",
        "Kora":          "kora kmena",
        "Nezdrava_hrca": "nezdrava hrca, poskodena tkanina",
        "Okolie":        "pozadie / okolie kmena",
        "Prasklina":     "trhliny a praskliny v dreve",
    }
    for name, desc in class_desc.items():
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(name)
        run.bold = True
        p.add_run(f" — {desc}")
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
# Section: 2. Methodology
# ═══════════════════════════════════════════════════════════════════════
def section_methodology(doc):
    add_heading(doc, "2  Metodika", 1)

    # ── 2.1 Dataset ──
    add_heading(doc, "2.1  Konstrukcia datasetu", 2)
    add_paragraph(doc,
        "Trenovaci dataset bol zostavený z 23 zaznamenanych a anotovanych "
        "dubovych kmenov, pricom kazdy kmen je "
        "reprezentovany niekolko desiatkami konsekutivnych priečnych "
        "rezov. Anotacia bola vykonana v MATLAB Image Labeler s "
        "pixelovou granularitou — kazda snimka ma asociovanu masku, kde "
        "kazdemu pixelu je priradena trieda. Pôvodne identifikatory "
        "(MATLAB konvencia 1–5) boli pri konverzii do Python prostredia "
        "premapovane na kanonicke poradie 0–4 (alphabeticke poradie "
        "tried) cez statickú look-up tabulku."
    )
    add_paragraph(doc, "Datasety boli rozdelene do styroch skupin:", bold=True)
    add_dataset_summary_table(doc)
    add_caption(doc, "Tabulka 1: Prehlad pouzitych datasetov a ich charakteristik.")
    doc.add_paragraph()
    add_paragraph(doc,
        "Datasety Dub_praskliny_a a hrce_mixed boli pridane v neskorsej "
        "faze projektu so zameranim na cielene zlepsenie detekcie vzacnych "
        "tried. Dub_praskliny_a obsahuje 50 ručne vybranych dubovych "
        "rezov, kde sa Prasklina vyskytuje v kazdej snimke (oproti "
        "~3,5 % v zvysku datasetu). Toto cielene rozsirenie tvori klucovu "
        "metodologickou volbu, ktorej prinos je kvantifikovany v sekcii 3.2."
    )
    doc.add_paragraph()

    # ── 2.2 Splits ──
    add_heading(doc, "2.2  Rozdelenie na trenovaciu, validacnu a "
                     "testovaciu mnozinu", 2)
    add_paragraph(doc,
        "Spravne navrhnute rozdelenie datasetu je v segmentacnych ulohach "
        "kriticke kvoli tzv. data leakage — situacii, ked susedne CT rezy "
        "rovnakeho kmena su priestorovo a vizualne velmi podobne. Pri "
        "naivnom nahodnom rozdelovani sa do trenovacej aj testovacej "
        "mnoziny dostanu rezy z toho isteho jedinca, co umelo nafukuje "
        "metricky. Naopak, prísne kmen-level rozdelenie (kazdy kmen patri "
        "iba do jednej mnoziny) odstrani leakage, ale v dosledku malej "
        "diverzity jedincov posunie validacne metrik nadol a sposobuje "
        "vacsi sum v early-stopping signali."
    )
    add_paragraph(doc,
        "V tejto praci sme zvolili hybridny prístup, ktory kombinuje vyhody "
        "oboch metod:"
    )
    for line in [
        "Testovacia mnozina (222 snimok, ~14,9 % datasetu) je definovana "
          "trunk-level: tri cele kmene (kmen4, kmen9, Dub_3b) plus prvych "
          "15 snimok z hrce_mixed a Dub_praskliny_a. Tieto jedince model "
          "nikdy nevidi pocas trénovania ani validacie, takze test "
          "predstavuje skutocnu generalizaciu na neznameho jedinca.",
        "Trenovacia (1015 snimok) a validacna (254 snimok) mnozina su "
          "vytvorene nahodnym 80/20 rozdelenim zvysneho pool-u (1175 snimok) "
          "bez ohladu na kmen. Dany within-trunk leakage v train/val sice "
          "nadhodnocuje validacnu mIoU, ale poskytuje hladsi signal pre "
          "early stopping a tuning hyperparametrov bez ovplyvnenia "
          "finalneho test vyhodnotenia.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    doc.add_paragraph()
    add_paragraph(doc,
        "Pri vybere testovacich kmenov bol zohladneny aspekt vzacnych tried — "
        "kombinacia kmen4 + kmen9 + Dub_3b zabezpecuje, ze testovacia "
        "mnozina obsahuje dostatocne pocet snimok s Prasklinou (216 z 222) "
        "a Nezdravou hrcou (57 z 222) pre statisticky robustny vypocet "
        "per-class metrik, hoci za cenu nadpriemerne hustoty Praskliny "
        "v teste (pomer Test%/Dataset% = 2,65). Tato „stress-test\" "
        "konfigurácia bola zámerná: pri trenovacej domanan pomalej "
        "konvergencii vzacnych tried je dolezite mat na teste dostatok "
        "pixelov pre stabilne odhady IoU."
    )
    doc.add_paragraph()
    add_paragraph(doc, "Tabulka 2 zobrazuje per-class pixelove rozdelenie "
                       "vo vsetkych troch mnozinach:", bold=False)
    doc.add_paragraph()
    add_split_distribution_table(doc)
    add_caption(doc, "Tabulka 2: Per-class pixelove rozdelenie (% pixelov) "
                     "v testovacej, trenovacej a validacnej mnozine.")
    doc.add_paragraph()

    # ── 2.3 Architecture ──
    add_heading(doc, "2.3  Architektura modelu", 2)
    add_paragraph(doc,
        "Pouzity bol model SegFormer-B2 [Xie 2021], variant strednej "
        "velkosti rodiny SegFormer s priblizne 25 milionmi parametrov. "
        "Architektura pozostáva z hierarchickeho transformer enkodera "
        "(Mix Transformer, MiT-B2), ktory generuje featurove mapy v "
        "stvoroch rozliseniach (1/4, 1/8, 1/16, 1/32 vstupnej rozlisenia), "
        "a z lahkeho MLP dekodera, ktory tieto featurove mapy fuzuje a "
        "produkuje finalnu segmentacnu mapu."
    )
    add_paragraph(doc,
        "Enkoder bol inicializovany predtrenovanymi vahami z ImageNet-1K "
        "klasifikacnej ulohy [Deng 2009]. Tento transfer learning sa v "
        "experimentoch ukazal ako kriticky — dva pokusy o 2.5D rozsirenie, "
        "ktore vyzadovali adaptaciu prvej konvolucnej vrstvy z 3 na 5 "
        "vstupnych kanalov a tym znicili predtrenovany signal, doasiahli "
        "horsie vysledky nez 2D baseline (sekcia 3.3). Vstupna rozlisenia "
        "modelu je 512×512 pixelov, vstupy su normalizovane ImageNet "
        "statistikami (mean = [0,485, 0,456, 0,406], std = [0,229, 0,224, 0,225]). "
        "Vystupom modelu su logity pre 5 tried pre kazdy pixel; finalna "
        "segmentacia sa získava operáciou argmax pozdĺz triednej osi."
    )
    doc.add_paragraph()

    # ── 2.4 Hyperparameters ──
    add_heading(doc, "2.4  Vyber hyperparametrov", 2)
    add_paragraph(doc,
        "Klucove hyperparametre boli volene na zaklade systematickych "
        "experimentov a iteracii. Najvyznamnejsim z nich je pomer vahy "
        "krížovej entropie (CE) a Dice straty v kombinovanej strate, "
        "ktoreho hodnota bola urcená cez 5-konfiguracny sweep:"
    )
    doc.add_paragraph()
    add_sweep_table(doc)
    add_caption(doc, "Tabulka 3: Vysledky sweep-u pomeru CE/Dice (5 konfiguracii). "
                     "Pomer 0,30/0,70 dosiahol najvyssie mIoU (75,19 %) a IoU "
                     "Praskliny (36,4 %).")
    doc.add_paragraph()
    add_paragraph(doc,
        "Vsetkych 5 konfiguracii bolo trenovanych s identickymi inymi "
        "hyperparametrami a identickym datasetom, aby bolo porovnanie "
        "metodologicky ciste. Vyhercom je pomer 0,30 CE / 0,70 Dice, "
        "ktory v porovnani s povodnym 0,35/0,65 baseline-om priniesol "
        "+1,14 percentuálneho bodu mIoU a +4,1 pp na IoU Praskliny. "
        "Tento pomer bol pouzity vo vsetkych nasledujucich experimentoch."
    )
    add_paragraph(doc,
        "Dalsie systematicky overené hyperparametre:"
    )
    for line in [
        "Learning rate: 6×10⁻⁵ (zvolene po porovnani 3e-5, 6e-5, 1e-4 — "
          "vyssia LR sposobovala nestabilnu konvergenciu pri Dice-heavy "
          "strate, nizsia spomalovala konvergenciu).",
        "Optimizer: AdamW (Loshchilov & Hutter, 2019) s weight_decay = "
          "10⁻⁴, ktory poskytuje lepsiu regularizaciu nez tradicny Adam.",
        "Scheduler: ReduceLROnPlateau s factorom 0,5 a patience 5, "
          "ktory automaticky znizuje LR pri stagnacii. Alternativa "
          "CosineAnnealingLR bola otestovana ale dosiahla horsie vysledky "
          "(73,30 % vs 73,83 %), pretoze cosine decay bol pre tento "
          "dataset prilis agresivny.",
        "Velkost batchu: 8 (limit GPU pamati pri rozliseni 512×512).",
        "Pocet epoch: max 200 s patience 25 — vyssia patience bola "
          "potrebna kvoli pomalsej konvergencii pri Dice-heavy strate "
          "(model konverguje typicky v 50–80 epochach).",
        "Class weights v strate: inverzne ku frekvencii triedy s clipom "
          "20 (predchadzanie extrémne vysokym vaham na vzacnych triedach).",
        "Oversample weighted random sampler: snimky s Prasklinou "
          "vzorkovane ×6, snimky s Nezdravou hrcou ×3 (uplna rovnomernost "
          "pri tak vzacnych triedach by sposobila degenerativny tréning).",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    doc.add_paragraph()

    # ── 2.5 Loss ──
    add_heading(doc, "2.5  Stratova funkcia", 2)
    add_paragraph(doc,
        "Pre hlavny model bola pouzita kombinovana strata zlozena z "
        "vázenej krížovej entropie a Dice straty:"
    )
    add_formula(doc,
        "L = 0,30 · CE(y, ŷ) + 0,70 · DiceLoss(y, ŷ)",
        "kde y je ground-truth maska a ŷ je predikcia modelu"
    )
    add_paragraph(doc,
        "Krížová entropia poskytuje stabilný gradient na pixelovej úrovni "
        "a explicitne penalizuje nespravne klasifikovane pixely. Dice "
        "strata [Milletari 2016] sa naopak zameriava na overlap "
        "predikovanej a skutocnej masky a je menej citliva na nevyvazenost "
        "tried. Ich kombinacia kompenzuje individualne slabiny — CE "
        "stabilita + Dice triedna citlivost. Pomer 0,30/0,70 bol "
        "empiricky urceny ako optimum (sekcia 2.4)."
    )
    add_paragraph(doc, "Per-class vahy v oboch komponentoch (ω_c) su pocítane "
                       "metodikou inverznej frekvencie:", bold=False)
    add_formula(doc,
        "ω_c = clamp(median(f) / f_c, max=20)",
        "kde f_c je frekvencia triedy c v trenovacom datasete"
    )
    add_paragraph(doc,
        "Pre pre experiment recall-optimalizacie (sekcia 3.4) bola Dice "
        "strata nahradena Focal-Tverskou stratou [Salehi 2017; Abraham 2019], "
        "ktora cez asymetricke parametre α (FP penalty) a β (FN penalty) "
        "umoznuje cielene riadit recall/precision trade-off:"
    )
    add_formula(doc,
        "TI(c) = TP / (TP + α(c)·FP + β(c)·FN)",
        "Tversky index pre triedu c"
    )
    add_formula(doc,
        "FT(c) = (1 − TI(c))^γ",
        "Focal-Tversky strata, γ = 1,33"
    )
    add_paragraph(doc,
        "Pri α(c) = β(c) = 0,5 sa Tversky index redukuje na klasicky Dice. "
        "Pre β > α model penalizuje FN viac nez FP, co zvysuje recall. "
        "Per-class hodnoty α a β su konkrétne diskutovane v sekcii 3.4."
    )
    doc.add_paragraph()

    # ── 2.6 Training ──
    add_heading(doc, "2.6  Trenovaci postup", 2)
    add_paragraph(doc,
        "Tréning prebiehal na jednom NVIDIA GPU s 12 GB VRAM (RTX-class). "
        "Pamatovo sme optimalizovali pouzitim AMP (automatic mixed "
        "precision) pre fp16 vypocty, gradient clippingom s max_norm = 2,0 "
        "(ochrana pred ojedinelymi explozivnymi gradientmi pri "
        "Dice/Focal-Tversky stratach), a presetreckov vsetkych vstupnych "
        "obrazkov a maskov v RAM (priemerne ~5 GB, dataset "
        "obsahuje ~1500 obrazkov)."
    )
    add_paragraph(doc, "Augmentacie aplikovane na trenovacie data:", bold=False)
    for line in [
        "HorizontalFlip s pravdepodobnostou 0,5",
        "VerticalFlip s pravdepodobnostou 0,5",
        "Rotate s rozsahom ±180° (CT rezy maju radialnu symetriu, plne "
          "rotacie su semanticky platne). Pre vyplnenie vznikajuceho "
          "okolia pri rotacii sa pri obrazku pouziva 0 a pri maske trieda "
          "Okolie (ID 3) — toto je klucovy detail, jeho oprava priniesla "
          "viditelne zlepsenie",
        "Resize na 512×512 a normalizacia ImageNet statistikami",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    add_paragraph(doc,
        "Validacne data prechadzaju iba operaciami Resize a Normalize. "
        "Tréningu predchadza zafixovanie nahodnych generatorov "
        "(seed = 42 pre Python random, NumPy a PyTorch CUDA) pre "
        "reprodukovatelnost. Vahy s najvyssou validacnou mIoU sa "
        "ukladaju a sluzia ako finalny model pre dany experiment."
    )
    doc.add_paragraph()

    # ── 2.7 Metrics ──
    add_heading(doc, "2.7  Hodnotiace metriky", 2)
    add_paragraph(doc,
        "Vsetky metriky su vypocítane z konfuznej matice C ∈ ℝ^(K×K) "
        "akumulovanej pocas inferencnej fázy na cele testovacie mnozine "
        "(K = 5 tried, jednotlive pixely sa pocítaju nezavisle). Pre triedu "
        "c sa definuju:"
    )
    bullet_items = [
        "TP_c (true positive)  — pocet pixelov spravne zaradenych do triedy c",
        "FP_c (false positive) — pocet pixelov inych tried nespravne "
          "predikovanych ako c",
        "FN_c (false negative) — pocet pixelov triedy c, ktore boli "
          "klasifikovane do inej triedy",
    ]
    for item in bullet_items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
    doc.add_paragraph()
    add_heading(doc, "2.7.1  IoU (Jaccardov index)", 3)
    add_paragraph(doc,
        "IoU (Intersection over Union) meria mieru prekryvu predikovanej "
        "a skutocnej masky. Hodnota 1 znamena dokonalu zhodu, 0 ziadny "
        "prekryv. Je to standardna metrika v semantickej segmentacii "
        "[Everingham 2015]."
    )
    add_formula(doc, "IoU(c) = TP_c / (TP_c + FP_c + FN_c)")
    add_formula(doc, "mIoU = (1/K) · Σ_c IoU(c)",
                "macro-priemer cez triedy (kazda trieda ma rovnaku váhu)")
    doc.add_paragraph()
    add_heading(doc, "2.7.2  Dice koeficient (F1 skore)", 3)
    add_paragraph(doc,
        "Dice koeficient je ekvivalentny F1 skore pre binarnu klasifikaciu "
        "kazdej triedy. V porovnani s IoU je linearnejsi pri malych "
        "objektoch. Pre kazdu triedu plati vzťah Dice = 2·IoU / (1 + IoU)."
    )
    add_formula(doc, "Dice(c) = 2·TP_c / (2·TP_c + FP_c + FN_c)")
    doc.add_paragraph()
    add_heading(doc, "2.7.3  Presnost (precision)", 3)
    add_paragraph(doc,
        "Presnost vyjadruje, akú cast pixelov predikovanych ako trieda c "
        "skutocne patri do triedy c. Nizka presnost znamena vela falosne "
        "pozitivnych predikcii (model „haluciuje\" triedu c)."
    )
    add_formula(doc, "Precision(c) = TP_c / (TP_c + FP_c)")
    doc.add_paragraph()
    add_heading(doc, "2.7.4  Uplnost (recall, sensitivity)", 3)
    add_paragraph(doc,
        "Uplnost vyjadruje, akú cast pixelov skutocne patriacich do "
        "triedy c bola modelom korektne detegovana. Pre vzacne triedy "
        "(Nezdrava_hrca, Prasklina) je recall klucova metrika v "
        "lesnickom kontexte: neidentifikovana vada (FN) je zavaznejsia "
        "chyba ako falosny poplach (FP), pretoze v drevarskej praxi je "
        "ekonomickejsie dodatocne overit oznacenu oblast nez nepostrehnut "
        "skrytu vadu kmena."
    )
    add_formula(doc, "Recall(c) = TP_c / (TP_c + FN_c)",
                "Synonyma: Sensitivity, True Positive Rate (TPR)")
    doc.add_paragraph()
    add_heading(doc, "2.7.5  Pixelova presnost", 3)
    add_paragraph(doc,
        "Celkovy podiel spravne klasifikovanych pixelov. Tato metrika je "
        "dominovana najfrekventovanejsimi triedami (v nasom datasete "
        "Okolie ~80 % pixelov), preto sa pouziva iba ako doplnkova k mIoU."
    )
    add_formula(doc, "PA = Σ_c TP_c / N",
                "kde N je celkovy pocet pixelov")
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
# Section: 3. Results
# ═══════════════════════════════════════════════════════════════════════
def section_results(doc, main, ablation, comp_3trunks, comp_3trunks_rgb,
                    ft_v1, ft_v2):
    add_heading(doc, "3  Vysledky", 1)

    # ── 3.1 Main model ──
    add_heading(doc, "3.1  Hlavny model (IoU-balanced)", 2)
    add_paragraph(doc,
        "Hlavny model (oznaceny v3) bol natrenovany na celej trenovacej "
        f"mnozine ({1015} snimok) so stratou 0,30·CE + 0,70·Dice po dobu "
        "77 epoch (early stopping pri patience = 25, najlepsia val mIoU "
        "v epoche 52). Tabulka 4 zobrazuje vysledky inference na "
        f"{main['n_images']}-snimkovej testovacej mnozine."
    )
    doc.add_paragraph()
    add_results_table(doc, main)
    add_caption(doc, "Tabulka 4: Per-class metriky hlavneho modelu (v3) "
                     "na testovacej mnozine.")
    doc.add_paragraph()

    miou = main["mean_iou"] * 100
    mrec = main["mean_recall"] * 100
    pa = main["pixel_accuracy"] * 100
    nez = main["per_class"]["Nezdrava_hrca"]
    prask = main["per_class"]["Prasklina"]

    add_paragraph(doc,
        f"Hlavny model dosiahol mIoU {miou:.2f} %, priemerny recall "
        f"{mrec:.2f} % a pixelovu presnost {pa:.2f} %. Dominantne triedy "
        f"(Drevo, Okolie) saturuju na hodnotach > 98 % IoU. Trieda Kora "
        f"({main['per_class']['Kora']['iou']*100:.1f} % IoU) ma vysoky "
        f"recall (97,9 %) ale nizsiu presnost — model spravne identifikuje "
        f"vsetky kôrove pixely, ale zaroven niektore okolite pixely "
        f"klasifikuje ako koru. Vzacne triedy: Nezdrava_hrca dosiahla "
        f"IoU {nez['iou']*100:.1f} %, recall {nez['recall']*100:.1f} %; "
        f"Prasklina IoU {prask['iou']*100:.1f} %, recall "
        f"{prask['recall']*100:.1f} %. Napriek extremnemu nepomerov tried "
        f"(0,02 % pixelov pre Prasklinu) model deteguje takmer 3/4 "
        f"vsetkych puklinovych pixelov."
    )
    doc.add_paragraph()

    # ── 3.2 Data ablation ──
    if ablation is not None:
        add_heading(doc, "3.2  Studia prispevku pridanych dat (data ablation)", 2)
        add_paragraph(doc,
            "Pre kvantifikaciu prinosu pridania cielene anotovanych dat "
            "(Dub_praskliny_a a hrce_mixed) bol natrenovany alternativny "
            "model na rovnakych hyperparametroch, ale s redukovanou "
            "trenovacou mnozinou (iba povodne kmen1–kmen10 a Dub_1–Dub_10). "
            "Test je identicky pre oba modely, takze ich rozdiel "
            "kvantifikuje cisty prinos pridanych dat."
        )
        doc.add_paragraph()
        add_comparison_table(doc, main, ablation,
                             label_main="v3 (s pridanymi datami)",
                             label_other="v3 ablation (bez pridanych)")
        add_caption(doc, "Tabulka 5: Vplyv pridanych datasetov na per-class "
                         "metriky. Pozitivne hodnoty Δ znamenaju, ze pridanie "
                         "dat zlepsilo metriku.")
        doc.add_paragraph()

        d_miou_a = (main["mean_iou"] - ablation["mean_iou"]) * 100
        prask_m = main["per_class"]["Prasklina"]
        prask_a = ablation["per_class"]["Prasklina"]
        nez_m = main["per_class"]["Nezdrava_hrca"]
        nez_a = ablation["per_class"]["Nezdrava_hrca"]
        d_rec_p = (prask_m["recall"] - prask_a["recall"]) * 100
        d_rec_n = (nez_m["recall"] - nez_a["recall"]) * 100

        add_paragraph(doc,
            f"Celkovy mIoU rozdiel je relatívne maly ({d_miou_a:+.2f} pp), "
            f"avsak detailny pohlad odhalí vyznamny posun v recall vzacnych "
            f"tried. Pridanim 50 cielene anotovanych snimok s Prasklinou "
            f"a podmnoziny hrce_mixed sa recall puklin zvysil o "
            f"{d_rec_p:+.1f} percentualneho bodu (z "
            f"{prask_a['recall']*100:.1f} % na {prask_m['recall']*100:.1f} %), "
            f"recall hrcí o {d_rec_n:+.1f} pp (z {nez_a['recall']*100:.1f} % "
            f"na {nez_m['recall']*100:.1f} %). Precision pri tom mierne "
            f"klesla — model je s pridanymi datami citlivejsi a deteguje "
            f"viac instancii vzacnych tried za cenu zvyseneho poctu false "
            f"positives. V kontexte detekcie poskodeni dreva, kde nezistena "
            f"vada je zavaznejsi error nez falosny poplach, je to ziaduca "
            f"zmiena."
        )
        add_paragraph(doc,
            "Z metodologickeho hladiska tento vysledok demonstrouje, ze "
            "agregovane metriky (mIoU) mozu zakryvat skutocnu povahu "
            "zlepsenia. Pre vzacne triedy je nutne sledovat per-class "
            "recall a precision oddelene, pretoze ich zmeny sa v mIoU "
            "navzajom kompenzuju."
        )
        doc.add_paragraph()

    # ── 3.3 2.5D experiments ──
    if comp_3trunks is not None:
        add_heading(doc, "3.3  Experimenty s 2.5D vstupom", 2)
        add_paragraph(doc,
            "Pre overenie, ci priestorova kontinuita medzi konsekutivnymi "
            "CT rezmi moze pomoct pri detekcii pukliín — ktorych "
            "geometricky charakter je inheretne 3D — boli implementovane "
            "dva alternativne 2.5D pristupy. V oboch pripadoch je vstupom "
            "modelu nie jeden CT rez, ale stack viacerych konsekutivnych "
            "rezov toho isteho kmena."
        )
        add_paragraph(doc,
            "Pre fer porovnanie 2D vs 2.5D bola vytvorena specialna "
            "testovacia mnozina obsahujuca iba 3 cele hold-out kmene "
            "(kmen4 + kmen9 + Dub_3b, 192 alebo 186 snimok podla varianty), "
            "kde su konsekutivne CT rezy garantovane priestorovo susediace. "
            "Snimky z Dub_praskliny_a (rucne vybrane, nie spatial neighbors) "
            "a hrce_mixed[:15] (out-of-distribution pre niektory typ "
            "trénovania) boli vylúčené."
        )
        doc.add_paragraph()

        # 3.3.1 5-channel
        m_v3_3t = comp_3trunks["v3"]
        m_abl_3t = comp_3trunks["v3_ablation"]
        m_25d_3t = comp_3trunks["v3_ablation_25d"]
        add_heading(doc, "3.3.1  5-kanaova varianta", 3)
        add_paragraph(doc,
            "Prvy variant pouzival 5 konsekutivnych CT rezov ako 5 vstupnych "
            "kanalov modelu. Tento pristup vyzaduje adaptaciu prvej "
            "konvolucnej vrstvy (in_channels = 3 → 5), priemerovanim "
            "RGB vah z ImageNet predtreningu. Ostatne hyperparametre "
            "ostali identicke s 2D ablation modelom."
        )
        doc.add_paragraph()
        add_3way_comparison_table(doc, m_v3_3t, m_abl_3t, m_25d_3t,
                                   labels=["v3 (2D, full)", "v3_abl (2D)",
                                           "v3_abl (2.5D)", "Δ 2.5D vs 2D"])
        add_caption(doc, "Tabulka 6: Porovnanie 2D a 5-kanalovej 2.5D varianty "
                         "na 192-snimkovej 3-trunks testovacej mnozine.")
        doc.add_paragraph()
        d_miou_25d = (m_25d_3t["mean_iou"] - m_abl_3t["mean_iou"]) * 100
        add_paragraph(doc,
            f"5-kanalovy 2.5D variant dosiahol o {abs(d_miou_25d):.2f} pp "
            f"nizsie mIoU nez ekvivalentny 2D model "
            f"({m_25d_3t['mean_iou']*100:.2f} % vs {m_abl_3t['mean_iou']*100:.2f} %). "
            "Hlavnym dovodom je strata ImageNet predtreningoveho signalu — "
            "adaptacia conv1 z 3 na 5 kanalov priemerovanim znicí kvalitu "
            "inicializacie."
        )
        doc.add_paragraph()

    # 3.3.2 RGB encoding
    if comp_3trunks_rgb is not None:
        m_2d_rgb = comp_3trunks_rgb["v3_ablation"]
        m_25d_rgb = comp_3trunks_rgb["v3_ablation_25d"]
        m_rgb = comp_3trunks_rgb["v3_ablation_rgb"]
        add_heading(doc, "3.3.2  RGB adjacent-slice encoding", 3)
        add_paragraph(doc,
            "Aby sme oddelili dva potencialne dovody zlyhania prvej "
            "varianty (strata ImageNet predtreningu vs. nedostatocny "
            "informacny obsah susednych rezov), implementovali sme "
            "alternativny pristup popularny v medicinskej obrazovke "
            "[Roth 2014; Christ 2016; Avesta 2023]: tri konsekutivne rezy "
            "su zakódovane ako jednotlive kanály RGB obrazka — R = (n−1), "
            "G = (n), B = (n+1). Tym sa zachovava 3-kanalova struktura a "
            "ImageNet predtrening funguje natívne. Centralny rez je v "
            "G kanali, ktory ma v luminosity formulácii (0,587·G) "
            "najvyssi vplyv na low-level features."
        )
        add_paragraph(doc,
            "Z trénovaciej aj testovacej mnoziny bol vyhodeny prvy a "
            "posledny rez kazdeho kmena (nemaju obojstrannych susedov). "
            "Test set sa redukoval na 186 snimok."
        )
        doc.add_paragraph()
        add_3way_comparison_table(doc, m_2d_rgb, m_25d_rgb, m_rgb,
                                   labels=["2D ablation", "2.5D 5-kanal",
                                           "RGB encoding", "Δ RGB vs 2D"])
        add_caption(doc, "Tabulka 7: Porovnanie 2D, 5-kanaloveho 2.5D a RGB "
                         "adjacent-slice encodingu na 186-snimkovej redukovanej "
                         "testovacej mnozine.")
        doc.add_paragraph()
        d_miou_rgb_2d = (m_rgb["mean_iou"] - m_2d_rgb["mean_iou"]) * 100
        d_miou_rgb_25d = (m_rgb["mean_iou"] - m_25d_rgb["mean_iou"]) * 100
        add_paragraph(doc,
            f"RGB encoding prekonal naivnu 5-kanalovu variantu o "
            f"{d_miou_rgb_25d:+.2f} pp mIoU, co potvrdzuje hypotezu o "
            f"dolezitosti zachovania 3-kanalovej struktury pre vyuzitie "
            f"ImageNet predtreningu. Avsak ani RGB encoding nedosiahol "
            f"vykon 2D baseline-u ({d_miou_rgb_2d:+.2f} pp). "
            f"Konzistentne zlyhanie dvoch nezavislych 2.5D pristupov "
            f"naznacuje, ze priestorova informacia zo susednych CT rezov "
            f"v dataste prierezov stromov nie je modelom dostatocne "
            f"vyuzitelna. Pravdepodobne pricinami su: (i) susedne CT "
            f"rezy stromov su anatomicky velmi podobne — neprínasaju "
            f"dostatok noveho informacneho obsahu; (ii) anotacie su "
            f"vylucne 2D, takze model sa supervizuje vzdy iba pre "
            f"stredny rez; (iii) pre G kanal RGB modelu vidi presne "
            f"ten isty vstup ako 2D model, dodatocna informacia v R/B "
            f"kanaloch je z perspektivy modelu skor sum nez signál."
        )
        doc.add_paragraph()

    # ── 3.4 Focal-Tversky ──
    if ft_v1 is not None and ft_v2 is not None:
        add_heading(doc, "3.4  Recall-optimalizovany model (Focal-Tversky)", 2)
        add_paragraph(doc,
            "V predchadzajucich vysledkoch hlavny model dosahoval recall "
            f"Praskliny {prask['recall']*100:.1f} % a recall Nezdravej hrcy "
            f"{nez['recall']*100:.1f} %. Pre praktické nasadenie modelu "
            "v lesnickom priemysle, kde nezistena vada predstavuje "
            "zavaznejsiu chybu nez falosny poplach, by bolo zelane recall "
            "vzacnych tried este zvysit, hoci aj za cenu znizenia "
            "presnosti a IoU. Pre tento ucel sme implementovali "
            "Focal-Tversky stratu s per-class α/β parametrami, ktora "
            "umoznuje cielene riadit recall/precision trade-off."
        )
        add_paragraph(doc,
            "Boli natrenovane dva varianty s rastucou agresiviou:"
        )
        for line in [
            "Focal-Tversky v1: β_Prasklina = 0,85 (β/α = 5,7), "
              "β_Nezdrava = 0,70 (β/α = 2,3)",
            "Focal-Tversky v2: β_Prasklina = 0,90 (β/α = 9,0), "
              "β_Nezdrava = 0,80 (β/α = 4,0)",
            "Triedy Drevo, Kora, Okolie maju α = β = 0,5 (Tversky "
              "redukovany na Dice — bez recall posunu)",
        ]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line)
        doc.add_paragraph()
        add_paragraph(doc,
            "Ostatne hyperparametre (LR, scheduler, augmentacie, sampler, "
            "class weights) ostali identicke s hlavnym modelom v3. "
            "Tabulka 8 zobrazuje progresiu vsetkych troch konfiguracii "
            "na rovnakej testovacej mnozine."
        )
        doc.add_paragraph()
        add_3way_comparison_table(doc, main, ft_v1, ft_v2,
                                   labels=["v3 (Dice baseline)",
                                           "FT v1 (β=0,85)", "FT v2 (β=0,90)",
                                           "Δ FT v2 vs v3"])
        add_caption(doc, "Tabulka 8: Vplyv Focal-Tverskej straty na recall "
                         "vzacnych tried. Per-class β parameter umoznuje "
                         "cielene zvysenie recall za kontrolovanu cenu IoU.")
        doc.add_paragraph()

        prask_v3 = main["per_class"]["Prasklina"]
        prask_ft1 = ft_v1["per_class"]["Prasklina"]
        prask_ft2 = ft_v2["per_class"]["Prasklina"]
        nez_v3 = main["per_class"]["Nezdrava_hrca"]
        nez_ft1 = ft_v1["per_class"]["Nezdrava_hrca"]
        nez_ft2 = ft_v2["per_class"]["Nezdrava_hrca"]
        d_prask_rec = (prask_ft2["recall"] - prask_v3["recall"]) * 100
        d_prask_iou = (prask_ft2["iou"] - prask_v3["iou"]) * 100
        d_nez_rec = (nez_ft2["recall"] - nez_v3["recall"]) * 100
        d_nez_iou = (nez_ft2["iou"] - nez_v3["iou"]) * 100

        add_paragraph(doc,
            f"Vysledky potvrdzuju hypotezu: pri agresivnejsich beta "
            f"hodnotach (FT v2) sa recall Praskliny zvysil o "
            f"{d_prask_rec:+.1f} pp (z {prask_v3['recall']*100:.1f} % na "
            f"{prask_ft2['recall']*100:.1f} %), za cenu zniženia IoU "
            f"o {d_prask_iou:+.1f} pp a precision o "
            f"{(prask_ft2['precision']-prask_v3['precision'])*100:+.1f} pp. "
            f"Pre Nezdravu hrcu sa recall zvysil o {d_nez_rec:+.1f} pp "
            f"({nez_v3['recall']*100:.1f} % → {nez_ft2['recall']*100:.1f} %) "
            f"pri zachovani IoU ({nez_v3['iou']*100:.1f} % → "
            f"{nez_ft2['iou']*100:.1f} %, Δ = {d_nez_iou:+.1f} pp) — "
            f"pri Nezdravej hrci β = 0,80 bolo dostatocne vyvazene "
            f"nastavenie, ktore poskytlo zlepsenie recall bez zhorsenia "
            f"prekryvu masky."
        )
        add_paragraph(doc,
            f"Trojica meraní (v3, FT v1, FT v2) tvori monotonný Pareto "
            f"front recall vs IoU pre Prasklinu, kde kazdy bod predstavuje "
            f"legitímne pracovne nastavenie modelu. Vyber medzi nimi "
            f"zavisi od konkrétneho aplikacneho scenara: balansovane "
            f"nasadenie (priemernu IoU klucova) preferuje v3, "
            f"recall-priorita preferuje FT v2."
        )
        doc.add_paragraph()
    elif ft_v2 is not None:
        # Iba FT v2 dostupny
        add_heading(doc, "3.4  Recall-optimalizovany model (Focal-Tversky)", 2)
        add_paragraph(doc,
            f"Pre maximalizaciu recall vzacnych tried bola pouzita "
            f"Focal-Tversky strata. Detailnejsie porovnanie viacerych "
            f"konfiguracii nie je v tejto verzii reportu zachytene; "
            f"finalny FT v2 model dosiahol Prasklina recall "
            f"{ft_v2['per_class']['Prasklina']['recall']*100:.1f} % "
            f"a Nezdrava recall {ft_v2['per_class']['Nezdrava_hrca']['recall']*100:.1f} %."
        )
        doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
# Section: 3.5 — v4 (rozsireny dataset) porovnanie
# ═══════════════════════════════════════════════════════════════════════
def section_v4_comparison(doc, v3_main, v4_main, v3_ft1, v4_ft1, v3_ft2, v4_ft2):
    add_heading(doc, "3.5  Vplyv dalsieho rozsirenia trenovacich dat (v4)", 2)
    add_paragraph(doc,
        "Po pociatocnom rozsireni datasetu o Dub_praskliny_a (50 snimok, "
        "studovane v sekcii 3.2) bol pridany aj komplementarny dataset "
        "Dub_praskliny_b (dalsich 50 snimok, sluburu '100 ks Dub praskliny_b', "
        "snimky 051–100). Spolu tak Dub_praskliny_a a Dub_praskliny_b "
        "tvoria povodnu 100-kusovu kolekciu rucne vybranych dubovych "
        "rezov so zameranim na detekciu pukliín."
    )
    add_paragraph(doc,
        "Vsetky predoslé výsledky (sekcie 3.1–3.4, oznacene v3) boli "
        "ziskane na trenovacej mnozine pred pridanim Dub_praskliny_b "
        "(1015 train, 254 val snimok). Pre kvantifikaciu prínosu daléhého "
        "rozsírenia boli rovnake architektury a hyperparametre znova "
        "natrenovane na rozsírenom datasete (1055 train, 264 val snimok, "
        "+50 snimok s vysokou hustotou Praskliny). Test set ostal "
        "identicky (222 snimok), takze v4 modely su priamo porovnatelne "
        "s v3."
    )
    doc.add_paragraph()

    # Tabulka 9: hlavny model v3 vs v4
    add_paragraph(doc,
        "Tabulka 9 zobrazuje vysledky hlavneho modelu (CE 30 % + Dice 70 %) "
        "pred a po rozsireni:", bold=False)
    doc.add_paragraph()
    add_comparison_table(doc, v4_main, v3_main,
                         label_main="v4 (1541 imgs)",
                         label_other="v3 (1491 imgs)")
    add_caption(doc, "Tabulka 9: Vplyv pridania Dub_praskliny_b na hlavny "
                     "model. Pozitivne Δ = v4 lepsi nez v3.")
    doc.add_paragraph()

    miou_d = (v4_main["mean_iou"] - v3_main["mean_iou"]) * 100
    p_v3 = v3_main["per_class"]["Prasklina"]
    p_v4 = v4_main["per_class"]["Prasklina"]
    add_paragraph(doc,
        f"Hlavny model dosiahol mIoU {v4_main['mean_iou']*100:.2f} % "
        f"({miou_d:+.2f} pp oproti v3, ktory mal {v3_main['mean_iou']*100:.2f} %). "
        f"Najvyznamnejsi posun bol pri triede Prasklina, kde IoU stupol "
        f"o {(p_v4['iou']-p_v3['iou'])*100:+.1f} pp "
        f"({p_v3['iou']*100:.1f} % → {p_v4['iou']*100:.1f} %) a precision "
        f"o {(p_v4['precision']-p_v3['precision'])*100:+.1f} pp. Recall sa "
        f"prakticky nezmenil ({p_v3['recall']*100:.1f} % → {p_v4['recall']*100:.1f} %), "
        f"co indikuje ze recall standarného modelu uz bol pri v3 v podstate "
        f"saturovany — pridane data prispeli k zvysenej presnosti predikcií, "
        f"nie k siroke detekcie."
    )
    doc.add_paragraph()

    # Tabulka 10: FT v2 v3 vs v4
    if v3_ft2 is not None and v4_ft2 is not None:
        add_paragraph(doc,
            "Tabulka 10 zobrazuje rovnake porovnanie pre recall-optimalizovany "
            "Focal-Tversky model (FT v2):", bold=False)
        doc.add_paragraph()
        add_comparison_table(doc, v4_ft2, v3_ft2,
                             label_main="v4 FT v2",
                             label_other="v3 FT v2")
        add_caption(doc, "Tabulka 10: Vplyv pridania Dub_praskliny_b na "
                         "recall-optimalizovany model.")
        doc.add_paragraph()

        p_ft_v3 = v3_ft2["per_class"]["Prasklina"]
        p_ft_v4 = v4_ft2["per_class"]["Prasklina"]
        d_rec = (p_ft_v4['recall'] - p_ft_v3['recall']) * 100
        d_iou = (p_ft_v4['iou'] - p_ft_v3['iou']) * 100
        add_paragraph(doc,
            f"Pri Focal-Tversky modeli sa naopak prejavil opačný efekt: "
            f"recall Praskliny stupol o {d_rec:+.1f} pp "
            f"(z {p_ft_v3['recall']*100:.1f} % na {p_ft_v4['recall']*100:.1f} %), "
            f"co prelomilo predoslý saturacny strop. IoU klesla iba "
            f"o {d_iou:+.1f} pp. Tento výsledok podporuje hypotezu, ze "
            f"pridane snimky s hustou Prasklinou pomohli modelu "
            f"identifikovat hraničné pripady (napr. tenké pukliny), ktore "
            f"baseline strata nedokazala efektivne sledovat."
        )
        doc.add_paragraph()

    # Pareto front summary
    add_paragraph(doc, "Aktualizovany Pareto front pre triedu Prasklina:",
                  bold=True)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["Model", "Loss", "IoU (%)", "Recall (%)", "Precision (%)"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
    shade_row(table.rows[0], "BDD7EE")

    # Build rows from available metrics
    rows_data = []
    if v3_main is not None:
        m = v3_main["per_class"]["Prasklina"]
        rows_data.append(("v3", "CE+Dice", m["iou"]*100, m["recall"]*100, m["precision"]*100))
    if v4_main is not None:
        m = v4_main["per_class"]["Prasklina"]
        rows_data.append(("v4", "CE+Dice", m["iou"]*100, m["recall"]*100, m["precision"]*100))
    if v3_ft1 is not None:
        m = v3_ft1["per_class"]["Prasklina"]
        rows_data.append(("v3 FT v1", "β=0,85", m["iou"]*100, m["recall"]*100, m["precision"]*100))
    if v4_ft1 is not None:
        m = v4_ft1["per_class"]["Prasklina"]
        rows_data.append(("v4 FT v1", "β=0,85", m["iou"]*100, m["recall"]*100, m["precision"]*100))
    if v3_ft2 is not None:
        m = v3_ft2["per_class"]["Prasklina"]
        rows_data.append(("v3 FT v2", "β=0,90", m["iou"]*100, m["recall"]*100, m["precision"]*100))
    if v4_ft2 is not None:
        m = v4_ft2["per_class"]["Prasklina"]
        rows_data.append(("v4 FT v2", "β=0,90", m["iou"]*100, m["recall"]*100, m["precision"]*100))

    for r in rows_data:
        row = table.add_row().cells
        row[0].text = r[0]
        row[1].text = r[1]
        row[2].text = f"{r[2]:.2f}"
        row[3].text = f"{r[3]:.2f}"
        row[4].text = f"{r[4]:.2f}"
    add_caption(doc, "Tabulka 11: Pareto front pre triedu Prasklina cez "
                     "vsetkych 6 trenovanych modelov. v4 modely posuvaju "
                     "front do vyssich hodnot ako IoU, tak recall.")
    doc.add_paragraph()

    add_paragraph(doc, "Diskusia rozsírenia v4:", bold=True)
    add_paragraph(doc,
        "Pridanie 50 dalsich snimok s hustou Prasklinou (Dub_praskliny_b) "
        "prinieslo viditelný ale uz mensí prinos nez prve rozsírenie "
        "o Dub_praskliny_a. Toto je v sulade s ocakavanou diminishing-returns "
        "krivkou: prve cielene anotovane data odstranili najvacsí deficit "
        "v reprezentacii vzacnej triedy, dalsie data prinasaju marginalnejsie "
        "zlepsenie pri rastucom celkovom objeme datasetu. Z metodologickeho "
        "hladiska je zaujimavy aj rozdielny charakter prinosu — pri "
        "balansovanom modeli (Dice) sa zlepsila precision a IoU, pri "
        "recall-orientovanom modeli (FT v2) sa zlepsil recall. To naznacuje, "
        "ze rozsirovanie datasetu posuva Pareto front, ale konkretny smer "
        "posunu zavisi od architektonickej volby loss funkcie. Finalny v4 "
        "FT v2 model dosahuje recall Praskliny "
        f"{v4_ft2['per_class']['Prasklina']['recall']*100:.1f} % a Nezdravej "
        f"hrcy {v4_ft2['per_class']['Nezdrava_hrca']['recall']*100:.1f} %, "
        f"co predstavuje doteraz najlepsie dosiahnute hodnoty pre tieto "
        f"klucove triedy v projektu."
    )
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
# Section: 4. Discussion
# ═══════════════════════════════════════════════════════════════════════
def section_discussion(doc, main, ft_v2):
    add_heading(doc, "4  Diskusia", 1)

    add_heading(doc, "4.1  Pareto trade-off recall vs IoU", 2)
    add_paragraph(doc,
        "Ako bolo demonstrovane v sekcii 3.4, recall vzacnych tried a "
        "IoU prekryv tvoria klasicky Pareto trade-off — zlepsenie jednej "
        "metriky vacsinou znamena zhorsenie druhej. Per-class β parameter "
        "v Tverskej strate poskytuje plynule riaditelnu „packu\" medzi "
        "tymito dvoma cielmi. Per-unit-β analyza ukazuje, ze marginalne "
        "vynasov recall klesaju pri vyssich hodnotach β: pri prechode "
        "z β = 0,5 na β = 0,85 sa Prasklina recall zvysil o 11,4 pp, "
        "ale zo β = 0,85 na β = 0,90 uz iba o 1,4 pp pri rovnakej "
        "absolutnej zmiene parametra (Δβ = 0,05). Toto naznacuje, ze "
        "Prasklina recall sa pri β ≥ 0,90 priblizuje saturacnemu stropu "
        "okolo 88 % na danej architekture a datasete."
    )
    add_paragraph(doc,
        "Pre triedu Nezdrava_hrca bol pozorovany odlisny trend: pri "
        "β = 0,80 sa recall zvysil o ~7 pp pri zachovani (dokonca miernom "
        "zvyseni) IoU. Toto naznacuje, ze pre Nezdravu hrcu sme este "
        "neboli na Pareto fronte v baseline modeli — model mal kapacitu "
        "zlepsit ako recall, tak presnost vyhladzovania masok zaroven. "
        "Tento jav sa interpretuje ako nasledok lepsej regularizacie cez "
        "dolne penalizovanie FN."
    )
    doc.add_paragraph()

    add_heading(doc, "4.2  Praktické implikacie pre nasadenie", 2)
    add_paragraph(doc,
        "Volba medzi modelmi v3 a FT v2 nie je „technicky lepsi vs horsi\", "
        "ale „rôzne pracovne body Pareto frontu pre rôzne aplikacie\":"
    )
    for line in [
        "v3 (IoU-balanced): vhodny pre aplikacie kde je dolezity presny "
          "tvar a hranice oblasti — napriklad meranie objemu vady, "
          "vizualizacia pre clovek-v-slucke kontrolu, kde sa pocita aj "
          "presnost predikcie. mIoU 79,5 % na teste.",
        "FT v2 (recall-optimized): vhodny pre prvotnu detekciu a screening — "
          "model „hlasi ked vidi cokolvek podozrele\" a odhalí 87 % "
          "vsetkych puklinovych pixelov a 91 % hrcovych pixelov. Vyssia "
          "miera FP (cca 60 % predikovanych puklinovych pixelov je "
          "v skutocnosti FP) je akceptovatelna za predpokladu, ze za "
          "modelom nasleduje kontrolnyj krok (manualna verifikacia, "
          "viacstupnova pipeline, atď.).",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    doc.add_paragraph()

    add_heading(doc, "4.3  Obmedzenia a buduca praca", 2)
    add_paragraph(doc,
        "Hlavne obmedzenia tejto prace:"
    )
    for line in [
        "Velkost a diverzita datasetu: vsetky data pochádzaju z duba — "
          "model nebol overovany na inych drevinach (smrek, buk, atd.). "
          "Buduca praca by mala overit prenositelnost modelu na ine "
          "druhy dreva a snimkovacie protokoly (ine CT skenery, rozlisenia).",
        "Trieda Prasklina ostava bottleneckom napriek vsetkym "
          "implementovanym strategiam (oversample, class weights, "
          "Focal-Tversky, rozsirenie datasetu). Maximum dosiahnutej IoU "
          "je 46,5 % (v4 baseline) a recall 88,7 % (v4 FT v2). Dalsie "
          "zlepsenie by pravdepodobne vyzadovalo zvysenie objemu cielene "
          "anotovanych puklin v trenovacom datasete (1–2 rády oproti "
          "súcasnemu) — krivka prinosov pri postupnom rozsirovani datasetu "
          "uz vykazuje znamky nasycenia (sekcia 3.5).",
        "Eksplicit 3D prístup nebol uspechom; potencial ma 3D U-Net "
          "s patch trenovanim, ktory si vsak vyzaduje 3D anotacie "
          "(zlucenie maskov pozdĺz Z-osi) a inu architekturu bez "
          "zavislosti na 2D ImageNet predtreningu.",
        "Test-time augmentation (TTA) a multi-scale inference neboli "
          "implementovane — su to overené techniky pre dodatocne "
          "zvysenie metrík na inferencnej strane bez zmien v trénovani.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
# Section: 5. Conclusion
# ═══════════════════════════════════════════════════════════════════════
def section_conclusion(doc, main, ft_v2, v4_main=None, v4_ft2=None):
    add_heading(doc, "5  Zaver", 1)
    # Pouzi v4 ak je k dispozicii, inak v3
    final_main = v4_main if v4_main is not None else main
    final_ft   = v4_ft2  if v4_ft2  is not None else ft_v2
    final_label = "v4" if v4_main is not None else "v3"

    miou = final_main["mean_iou"] * 100
    main_prask_rec = final_main["per_class"]["Prasklina"]["recall"] * 100
    main_nez_rec = final_main["per_class"]["Nezdrava_hrca"]["recall"] * 100
    add_paragraph(doc,
        f"V tejto praci bol vyvinuty a evalvovany system semantickej "
        f"segmentacie CT rezov dubovych kmenov zalozeny na architekture "
        f"SegFormer-B2. Finalny IoU-balanced model ({final_label}) trenovany "
        f"na kombinovanej strate 0,30·CE + 0,70·Dice dosiahol mIoU "
        f"{miou:.2f} %, recall Praskliny {main_prask_rec:.1f} % a recall "
        f"Nezdravej hrcy {main_nez_rec:.1f} % na hold-out testovacej mnozine."
    )
    if final_ft is not None:
        ft_prask = final_ft["per_class"]["Prasklina"]["recall"] * 100
        ft_nez = final_ft["per_class"]["Nezdrava_hrca"]["recall"] * 100
        ft_miou = final_ft["mean_iou"] * 100
        add_paragraph(doc,
            f"Pre praktické nasadenie v scenaroch s recall-prioritou bola "
            f"vyvinuta alternativa zalozena na Focal-Tverskej strate s "
            f"per-class β parametrami (FT v2), ktora dosiahla recall "
            f"Praskliny {ft_prask:.1f} % a recall Nezdravej hrcy "
            f"{ft_nez:.1f} % pri mIoU {ft_miou:.2f} %. Oba modely "
            f"predstavuju legitimne pracovne body Pareto frontu medzi "
            f"recall a IoU; vyber medzi nimi zavisi od konkrétnej "
            f"aplikácie."
        )
    add_paragraph(doc,
        "Klucove metodologicke prinosy tejto prace su: (i) demonstracia "
        "dolezitosti hybridneho train/val/test rozdelenia s trunk-level "
        "test mnozinou pre prevenciu data leakage; (ii) kvantifikacia "
        "prinosu cielenej anotacie vzacnych tried na priklade dvoch po sebe "
        "idúcich rozsírení datasetu (Dub_praskliny_a a Dub_praskliny_b) — "
        "viditelný posun Pareto frontu pri postupnom raste objemu "
        "anotovanych dat; (iii) overenie, ze naive 2.5D rozsirenie (ako "
        "5-kanalove tak RGB encoding) nie je v tomto domene uzitocné; "
        "(iv) demonstracia kontrolovatelneho recall/IoU trade-offu cez "
        "per-class Focal-Tversky parametre. Vsetky experimenty su plne "
        "reprodukovatelne — kod aj vahy modelov su zverejnene v projekte."
    )
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════
# Section: 6. References
# ═══════════════════════════════════════════════════════════════════════
def section_references(doc):
    add_heading(doc, "6  Pouzita literatura", 1)
    refs = [
        "Abraham, N., Khan, N. M. (2019). A Novel Focal Tversky Loss "
            "Function with Improved Attention U-Net for Lesion Segmentation. "
            "IEEE International Symposium on Biomedical Imaging (ISBI).",
        "Avesta, A., et al. (2023). Comparing 3D, 2.5D, and 2D Approaches "
            "to Brain Image Auto-Segmentation. Bioengineering, 10(2):181.",
        "Christ, P. F., et al. (2016). Automatic Liver and Lesion "
            "Segmentation in CT using Cascaded Fully Convolutional Neural "
            "Networks and 3D Conditional Random Fields. MICCAI.",
        "Deng, J., et al. (2009). ImageNet: A Large-Scale Hierarchical "
            "Image Database. IEEE Conference on Computer Vision and "
            "Pattern Recognition (CVPR).",
        "Everingham, M., et al. (2015). The Pascal Visual Object Classes "
            "Challenge: A Retrospective. International Journal of Computer "
            "Vision, 111(1):98–136.",
        "Loshchilov, I., Hutter, F. (2019). Decoupled Weight Decay "
            "Regularization. International Conference on Learning "
            "Representations (ICLR).",
        "Milletari, F., Navab, N., Ahmadi, S.-A. (2016). V-Net: Fully "
            "Convolutional Neural Networks for Volumetric Medical Image "
            "Segmentation. Fourth International Conference on 3D Vision "
            "(3DV).",
        "Roth, H. R., et al. (2014). A New 2.5D Representation for Lymph "
            "Node Detection using Random Sets of Deep Convolutional Neural "
            "Network Observations. MICCAI.",
        "Salehi, S. S. M., Erdogmus, D., Gholipour, A. (2017). Tversky "
            "Loss Function for Image Segmentation Using 3D Fully "
            "Convolutional Deep Networks. International Workshop on "
            "Machine Learning in Medical Imaging.",
        "Xie, E., et al. (2021). SegFormer: Simple and Efficient Design "
            "for Semantic Segmentation with Transformers. Advances in "
            "Neural Information Processing Systems (NeurIPS), 34.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(r)
        run.font.size = Pt(10)


# ═══════════════════════════════════════════════════════════════════════
# Build report
# ═══════════════════════════════════════════════════════════════════════
def build_report(doc, main, ablation, comp_3trunks, comp_3trunks_rgb,
                 ft_v1, ft_v2, v4_main=None, v4_ft1=None, v4_ft2=None):
    # Title
    title = doc.add_heading(
        "Semanticka segmentacia CT rezov dubovych kmenov pomocou SegFormer-B2: "
        "metodika, hyperparametricka studia a recall-optimalizovane modely", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    section_intro(doc)
    section_methodology(doc)
    section_results(doc, main, ablation, comp_3trunks, comp_3trunks_rgb,
                    ft_v1, ft_v2)

    # 3.5 v4 comparison (only if v4 main is loaded)
    if v4_main is not None:
        section_v4_comparison(doc, main, v4_main, ft_v1, v4_ft1, ft_v2, v4_ft2)

    section_discussion(doc, main, ft_v2)
    section_conclusion(doc, main, ft_v2, v4_main, v4_ft2)
    section_references(doc)


# ═══════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════
def load_metrics(path: Path, label: str):
    if path is None or not path.exists():
        print(f"  {label:<25}: NOT FOUND ({path})")
        return None
    with open(path) as f:
        m = json.load(f)
    print(f"  {label:<25}: {path.name} -> mIoU {m['mean_iou']*100:.2f}%")
    return m


def main(args):
    print("Loading metrics:")
    main_metrics      = load_metrics(args.main, "main (v3)")
    if main_metrics is None:
        print("ERROR: hlavny model metrics neexistuje")
        return
    ablation_metrics  = load_metrics(args.ablation, "ablation")

    comp_3trunks = None
    if args.v3_3trunks and args.ablation_3trunks and args.ablation_25d_3trunks:
        c = {
            "v3":              load_metrics(args.v3_3trunks,         "3trunks v3"),
            "v3_ablation":     load_metrics(args.ablation_3trunks,    "3trunks ablation"),
            "v3_ablation_25d": load_metrics(args.ablation_25d_3trunks, "3trunks 2.5D"),
        }
        if all(c.values()):
            comp_3trunks = c

    comp_3trunks_rgb = None
    if args.ablation_3trunks_rgb and args.ablation_25d_3trunks_rgb \
            and args.ablation_rgb_3trunks:
        c = {
            "v3_ablation":     load_metrics(args.ablation_3trunks_rgb,    "rgb 2D"),
            "v3_ablation_25d": load_metrics(args.ablation_25d_3trunks_rgb, "rgb 2.5D"),
            "v3_ablation_rgb": load_metrics(args.ablation_rgb_3trunks,    "rgb RGB"),
        }
        if all(c.values()):
            comp_3trunks_rgb = c

    ft_v1 = load_metrics(args.ft_v1, "FT v1")
    ft_v2 = load_metrics(args.ft_v2, "FT v2")

    # v4 (rozsireny dataset s Dub_praskliny_b)
    v4_main = load_metrics(args.v4_main, "v4 main")
    v4_ft1  = load_metrics(args.v4_ft_v1, "v4 FT v1")
    v4_ft2  = load_metrics(args.v4_ft_v2, "v4 FT v2")

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(12)

    build_report(doc, main_metrics, ablation_metrics, comp_3trunks,
                 comp_3trunks_rgb, ft_v1, ft_v2,
                 v4_main=v4_main, v4_ft1=v4_ft1, v4_ft2=v4_ft2)

    doc.save(args.out)
    print(f"\nReport saved -> {args.out}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--main",     type=Path,
                        default=DATA_DIR / "metrics_v3.json")
    parser.add_argument("--ablation", type=Path,
                        default=DATA_DIR / "metrics_v3_ablation.json")
    parser.add_argument("--v3-3trunks",          type=Path,
                        default=DATA_DIR / "metrics_v3_3trunks.json")
    parser.add_argument("--ablation-3trunks",    type=Path,
                        default=DATA_DIR / "metrics_v3_ablation_3trunks.json")
    parser.add_argument("--ablation-25d-3trunks", type=Path,
                        default=DATA_DIR / "metrics_v3_ablation_25d_3trunks.json")
    parser.add_argument("--ablation-3trunks-rgb", type=Path,
                        default=DATA_DIR / "metrics_v3_ablation_3trunks_rgb.json")
    parser.add_argument("--ablation-25d-3trunks-rgb", type=Path,
                        default=DATA_DIR / "metrics_v3_ablation_25d_3trunks_rgb.json")
    parser.add_argument("--ablation-rgb-3trunks", type=Path,
                        default=DATA_DIR / "metrics_v3_ablation_rgb_3trunks.json")
    parser.add_argument("--ft-v1", type=Path,
                        default=DATA_DIR / "metrics_v3_focaltversky.json")
    parser.add_argument("--ft-v2", type=Path,
                        default=DATA_DIR / "metrics_v3_focaltversky_v2.json")
    parser.add_argument("--v4-main", type=Path,
                        default=DATA_DIR / "metrics_v4.json")
    parser.add_argument("--v4-ft-v1", type=Path,
                        default=DATA_DIR / "metrics_v4_focaltversky.json")
    parser.add_argument("--v4-ft-v2", type=Path,
                        default=DATA_DIR / "metrics_v4_focaltversky_v2.json")
    parser.add_argument("--out",   type=Path, default=DATA_DIR / "report.docx")
    args = parser.parse_args()
    main(args)
